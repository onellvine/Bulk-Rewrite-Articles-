from docx import Document
import os
from selenium import webdriver
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.common.by import By
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import TimeoutException
from selenium.webdriver.support.ui import WebDriverWait
from pprint import pprint


def getFiles(fold):
    files = os.listdir(fold)
    file_names = []

    for file in files:
        if file.endswith('.docx') or file.endswith('doc'):
            file_names.append(os.path.splitext(file)[0])

    return file_names

def element_presence(by,xpath,time):
    element_present = EC.presence_of_element_located((By.XPATH, xpath))
    WebDriverWait(driver, time).until(element_present)


app = "https://litessay.com/"
driver = webdriver.Chrome("F:\\Documents\\web\\jss\\bulk rewrite\\chromedriver.exe")

doc_files = getFiles(os.getcwd())

for i in range(len(doc_files)):
    #opening the docx file containing text to be rewritten 
    document = Document('F:\\Documents\\web\\jss\\bulk rewrite\\{}.docx'.format(doc_files[i]))
    driver.get(app)

    input_area = driver.find_element(By.XPATH, '//*[@id="ogtxt"]')

    content = []

    for p in document.paragraphs:
        content.append(p.text.lstrip("\t"))
        content.append("\n")

    essay = ''.join(content)

    #populate original essay
    input_area.send_keys(essay)

    #press the rewrite button
    rewrite_button = driver.find_element(By.XPATH, '/html/body/div/div/div[2]/form/button')
    rewrite_button.click()

    #wait fro reload of page to rewrite the article
    element_presence(By.XPATH, '//*[@id="newtxt"]', 10)

    #get the rewritten essay
    new_essay = driver.find_element_by_id('newtxt').get_attribute('value')

    #create a new document file to save the rewritten essay
    new_book1 = Document()
    new_book1.add_paragraph(new_essay)
    new_book1.save('{}_rewritten.docx'.format(doc_files[i]))

#quit operation after all essays have been rewritten
driver.quit()

